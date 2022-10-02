from PIL import Image
import os, numpy as np
import docx2txt
import docx
import aspose.words as aw

MAX_COLOR_VALUE = 255
MAX_BIT_VALUE = 8

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

def to_bin(data):
    """Convert `data` to binary format as string"""
    if isinstance(data, str):
        return ''.join([ format(ord(i), "08b") for i in data ])
    elif isinstance(data, bytes) or isinstance(data, np.ndarray):
        return [ format(i, "08b") for i in data ]
    elif isinstance(data, int) or isinstance(data, np.uint8):
        return format(data, "08b")
    else:
        raise TypeError("Type not supported.")
        
# Creates an image with the RGB values
def createImage(data, resolution):
    image = Image.new("RGB", resolution)
    image.putdata(data)
    return image


def remove_n_least_significant_bits(value, n):
    value = value >> n
    return value << n


def get_n_least_significant_bits(value, n):
    value = value << MAX_BIT_VALUE - n
    value = value % MAX_COLOR_VALUE
    return value >> MAX_BIT_VALUE - n


def get_n_most_significant_bits(value, n):
    return value >> MAX_BIT_VALUE - n


def shift_n_bits_to_8(value, n):
    return value << MAX_BIT_VALUE - n

def saveDocxFile(data, filename):
    textList = ""
    encoded_doc = docx.Document()
    for i in data:
        textList += i
    
    #print(textList)
    cleaned_string = ''.join(c for c in textList if valid_xml_char_ordinal(c))
    encoded_doc.add_paragraph(cleaned_string)
    encoded_doc.save(filename)

def asciiListForTxtFile(filename, list):
    with open(filename, "rb") as out:
        while True:
            c = out.read(1)
            if not c:
                print("End of file")
                break
            list.append(ord(c))
    print(list)

def asciiListForDocxFile(filename, list):
    fullText = []
    doc = docx.Document(filename)
    for i in doc.paragraphs:
        #print(i.text)
        fullText.append(i.text)
    docText = '\n'.join(fullText)
    for i in docText:
        #print(ord(i))
        list.append(ord(i))

# Function to encode the image that we want to hide, into the cover image
def encodeTextFile(file_to_hide, file_to_hide_in, n_bits):
    toHide_file_type = file_to_hide.split(".")[1]
    print(toHide_file_type)
    
    #Opening of payload file
    # List to store the ascii values from the content of the file
    payload_bytes = []

    if(toHide_file_type == "txt"):
        #Uncomment to retrieve txt file text
        asciiListForTxtFile(file_to_hide, payload_bytes)
    elif(toHide_file_type == "docx"):
        #Uncomment to retrieve docx text
        asciiListForDocxFile(file_to_hide, payload_bytes)

    """fullText = []
    doc = docx.Document(file_to_hide)
    for i in doc.paragraphs:
        #print(i.text)
        fullText.append(i.text)
    docText = '\n'.join(fullText)
    for i in docText:
        #print(ord(i))
        payload_bytes.append(ord(i))"""

    print("ASD ", payload_bytes)

    #Opening of cover file
    cover_bytes = []
    with open(file_to_hide_in, "rb") as out:
        while True:
            c = out.read(1)
            if not c:
                print("End of file")
                break
            cover_bytes.append(ord(c))
    #print(cover_bytes)


    payload_size = os.path.getsize(file_to_hide)  # file_size is the same as len(file_bytes)
    print("The payload file size is ", payload_size)

    cover_size = os.path.getsize(file_to_hide_in)
    print("The cover file size is ", cover_size)

    # Check if cover file size is smaller than payload img size
    if (cover_size < payload_size):
        print("Cover object size is smaller than payload size")
        quit()

    data = []
    for i in range(len(payload_bytes)):
        char_hide = payload_bytes[i]
        char_hide = get_n_most_significant_bits(char_hide, n_bits)

        char_hide_in = cover_bytes[i]
        char_hide_in = remove_n_least_significant_bits(char_hide_in, n_bits)

        data.append(char_hide + char_hide_in)

    with open("encoded_file.docx", "w", encoding="utf-8") as f:
        for i in data:
            f.write(chr(i))
    #print(data)

    """#Part for encoding into docx file. NOT WORKING
    textList = ""
    encoded_doc = docx.Document()
    for i in data:
        textList += chr(i)
    
    print(textList)
    encoded_doc.add_paragraph(textList)
    encoded_doc.save("encoded_file.docx")"""
    
    return None


#Decode image and n_bits as parameters.
def decodeFile(file_to_decode, n_bits):
    asciiArray = []
    with open(file_to_decode, "r", encoding="utf-8") as f:
        for i in f.read():
            #print(i)
            asciiArray.append(ord(i))
    
    print(asciiArray)

    data = ""
    for i in range(len(asciiArray)):
        char_encoded = asciiArray[i]
        char_encoded = get_n_least_significant_bits(char_encoded, n_bits)

        char_encoded = shift_n_bits_to_8(char_encoded, n_bits)

        #print(chr(char_encoded))
        data += chr(char_encoded)

    print(data)

    with open("decoded_file.txt", "w", encoding="utf-8") as f:
        f.write(data)

    
    #saveDocxFile(data, "decoded_file.docx")

    return None

# # Running encode function
# # no. of bits
n_bits = 8
# # path of files replace them as per your need.
#txt to txt
#file_to_hide_path = "payload_file.txt"
#cover_file_path = "test.txt"

#txt to docx
file_to_hide_path = "payload_file.txt"
cover_file_path = "wordTest.docx"

#docx to txt
#file_to_hide_path = "payload_file_docx.docx"
#cover_file_path = "test.txt"
encodeTextFile(file_to_hide_path, cover_file_path, n_bits)
print("File encoded Successfully!")
#
# # running the decode function
n_bits = 8
#
# # path where you would want to save decoded Image.
#file_to_decode_path = "encoded_file.docx"

file_to_decode_path = "encoded_file.docx" #Change in line 149 too
decodeFile(file_to_decode_path, n_bits)
print("Image decoded Successfully!")
