from PIL import Image
import os, numpy as np
import docx2txt
import docx

MAX_COLOR_VALUE = 255
MAX_BIT_VALUE = 8

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

# Creates an image with the RGB values
def createImage(data, resolution):
    image = Image.new("RGB", resolution)
    image.putdata(data)
    return image


def remove_n_least_significant_bits(value, n):
    value = value >> n
    return value << n


def get_n_least_significant_bits(value, n):
    value = value << MAX_BIT_VALUE - n
    value = value % MAX_COLOR_VALUE
    return value >> MAX_BIT_VALUE - n


def get_n_most_significant_bits(value, n):
    return value >> MAX_BIT_VALUE - n


def shift_n_bits_to_8(value, n):
    return value << MAX_BIT_VALUE - n

def getHeightWidth(asciiArray):
    counter = 0
    index = 0
    heightWidth = None
    for i in range(len(asciiArray)):
        #print(asciiArray[i])
        if (counter == 5):
            index = i
            #print(i)
            break
            
        if (asciiArray[i] == 61):
            counter += 1
            #print("hi")
        else:
            counter = 0
    #print("buh ", index)
    newAsciiArray = asciiArray[:index]
    asciiHeightWidth = asciiArray[index:]
    heightWidth = []
    for i in asciiHeightWidth:
        heightWidth.append(chr(i))
    
    heightArray = []
    widthArray = []
    for i in range(len(heightWidth)):
        print(heightWidth[i])
        if (heightWidth[i] == "="):
            print(i)
            heightArray = heightWidth[i:]
            break
    heightArray = heightArray[1:]
    height2 = []
    for i in range(len(heightArray)):
        if (heightArray[i] == "W"):
            height2 = heightArray[:i]
            widthArray = heightArray[i:]
            break
    for i in range(len(widthArray)):
        if (widthArray[i] == "="):
            widthArray = widthArray[i:]
            break
    widthArray = widthArray[1:]

    height, width = "", ""
    for i in height2:
        height += i
    for i in widthArray:
        width += i

    return height,width,newAsciiArray


# Function to encode the image that we want to hide, into the cover image
def encodeImage(image_to_hide, file_to_hide_in, n_bits, outputFile):

    # List to store the ascii values from the content of the file
    file_bytes = []

    with open(file_to_hide_in, "rb") as out:
        while True:
            c = out.read(1)
            if not c:
                print("End of file")
                break
            file_bytes.append(ord(c))

    # Check size of cover file
    file_size = os.path.getsize(file_to_hide_in)  # file_size is the same as len(file_bytes)
    print("The txt file size is ", file_size)

    # Keep adding character till it can be reshaped into a list with 3 elements (In the form of RGB (1,2,3))
    while (len(file_bytes) % 3 != 0):
        file_bytes.append(ord("A"))

    file_bytes = np.array(file_bytes)
    file_bytes = np.reshape(file_bytes, (-1, 3))

    # Check size of payload img
    width = image_to_hide.size[0]
    height = image_to_hide.size[1]
    print("SIZE OF IMG ", image_to_hide.size)
    img_size = (width * height) * 3  # multiply by 3 cuz there's 3 bytes per pixel

    # Check if cover file size is smaller than payload img size
    if (file_size < img_size):
        print("Cover object size is smaller than payload size")
        quit()

    # .load() returns the "pixel_access" object that has the data(matrix) of the pixels.
    hide_image = image_to_hide.load()

    # this will store the values of each individual pixel as a matrix.
    data = []
    counter = 0  # Loop thru the txt file
    # looping the hide_image object.
    print("HEIGHT ", height)
    print("WIDTH ", width)
    for y in range(height):
        for x in range(width):
            try:
                # gets n most significant bits of r,g,b values of image to hide.
                # print("HIDE IMAGE ", hide_image[x,y])
                r_hide, g_hide, b_hide = hide_image[x, y]

                r_hide = get_n_most_significant_bits(r_hide, n_bits)
                g_hide = get_n_most_significant_bits(g_hide, n_bits)
                b_hide = get_n_most_significant_bits(b_hide, n_bits)

                # remove least n significant bits of txt file's content to hide in so we can store
                # the n most significant bits in that place.

                r_hide_in, g_hide_in, b_hide_in = file_bytes[counter]

                r_hide_in = remove_n_least_significant_bits(r_hide_in, n_bits)
                g_hide_in = remove_n_least_significant_bits(g_hide_in, n_bits)
                b_hide_in = remove_n_least_significant_bits(b_hide_in, n_bits)

                # Adding each MSB R,G,B from payload to LSB R,G,B cover image

                data.append((r_hide + r_hide_in,
                             g_hide + g_hide_in,
                             b_hide + b_hide_in))
                counter += 1

            # incase of exception it will show the reason.
            except Exception as e:
                print(e)

    # print("Data ", data)
    data = np.array(data)
    data = data.flatten()

    # Writing encoded characters to encoded_file.txt
    with open(outputFile, "w", encoding="utf-8") as f:
        for i in data:
            f.write(chr(i))
            # charData.append(chr(i))
        print(str(height), " ", str(width))
        f.write("=====")
        f.write("HEIGHT=")
        f.write(str(height))
        f.write("WIDTH=")
        f.write(str(width))


    #WRITING TO WORD DOC. NOT WORKING    
    """textList = []
    decoded_doc = docx.Document()
    for i in data:
        #print(chr(i))
        x = chr(i)
        #print(x)
        cleaned_string = ''.join(c for c in x if valid_xml_char_ordinal(c))
        #print(cleaned_string)
        textList.append(cleaned_string)
    
    fullText = ""
    for i in textList:
        fullText += i
    #print(fullText)
    fullText += ("=====")
    fullText += ("HEIGHT=")
    fullText += (str(height))
    fullText += ("WIDTH=")
    fullText += (str(width))
    decoded_doc.add_paragraph(fullText)
    decoded_doc.save("encoded_file.docx")"""

    # return an height and width from the above data.
    return height, width


# Decode image and n_bits as parameters.
def decodeFile(file_to_decode, n_bits):
    # List to store the ascii values from the content of the file
    asciiArray = []
    with open(file_to_decode, "r", encoding="utf-8") as f:
        for i in f.read():
            #print(i)
            asciiArray.append(ord(i))

    #READ WORD DOC. NOT WORKING
    """fullText = []
    doc = docx.Document(file_to_decode)
    for i in doc.paragraphs:
        #print(i.text)
        fullText.append(i.text)
    docText = '\n'.join(fullText)
    for i in docText:
        #print(ord(i))
        asciiArray.append(ord(i))"""
    
    #print("TEST ", len(asciiArray))

    height,width,newAsciiArray = getHeightWidth(asciiArray)
    
    height = int(height)
    width = int(width)
    print(height, " ", width)


    # Keep adding character till it can be reshaped into a list with 3 elements (In the form of RGB (1,2,3))
    while (len(newAsciiArray) % 3 != 0):
        newAsciiArray.append(ord("A"))
        # print(file_size % 3)
    
    #print(len(newAsciiArray))
    npAsciiArray = np.array(newAsciiArray)
    npAsciiArray = np.reshape(npAsciiArray, (-1, 3))
    #print("ASCIITEST ", len(npAsciiArray))

    # matrix that will store the extracted pixel values from the encoded txt file.
    data = []
    counter = 0
    # looping through every pixel in the encoded Image.
    for y in range(height):
        for x in range(width):
            # gets rgb values of encoded txt file.
            r_encoded, g_encoded, b_encoded = npAsciiArray[counter]

            # get n least significant bits for each r,g,b value of the encoded image
            r_encoded = get_n_least_significant_bits(r_encoded, n_bits)
            g_encoded = get_n_least_significant_bits(g_encoded, n_bits)
            b_encoded = get_n_least_significant_bits(b_encoded, n_bits)

            # shifts the n bits to right so that they occupy a total of 8 bit spaces.
            # If 10 are the bits then shifting them would look like 10000000
            # this would ofcourse be converted to an int as per python's bit operations.

            r_encoded = shift_n_bits_to_8(r_encoded, n_bits)
            g_encoded = shift_n_bits_to_8(g_encoded, n_bits)
            b_encoded = shift_n_bits_to_8(b_encoded, n_bits)

            data.append((r_encoded, g_encoded, b_encoded))
            counter += 1
            
    # print("DATA ", data)

    return createImage(data, (width, height))


# # Running encode function
# # no. of bits
n_bits = 2
# # path of files replace them as per your need.
image_to_hide_path = "tree2.jpg"
image_to_hide_in_path = "test.txt"
#image_to_hide_in_path = "wordTest.docx"
#image_to_hide_in_path = "excel.xlsx"
image_to_hide = Image.open(image_to_hide_path)
encodeImage(image_to_hide, image_to_hide_in_path, n_bits, "encoded_file.txt")
print("File encoded Successfully!")
#
# # running the decode function
n_bits = 2
#
# # path where you would want to save decoded Image.
decoded_image_path = "pics/decoded_file.png"
file_to_decode = "encoded_file.txt"
decodeFile(file_to_decode, n_bits).save(decoded_image_path)
print("Image decoded Successfully!")
