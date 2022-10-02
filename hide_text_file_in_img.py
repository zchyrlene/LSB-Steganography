from PIL import Image
import os, numpy as np
import docx2txt
import docx
import aspose.words as aw

MAX_COLOR_VALUE = 255
MAX_BIT_VALUE = 8

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

def to_bin(data):
    """Convert `data` to binary format as string"""
    if isinstance(data, str):
        return ''.join([ format(ord(i), "08b") for i in data ])
    elif isinstance(data, bytes) or isinstance(data, np.ndarray):
        return [ format(i, "08b") for i in data ]
    elif isinstance(data, int) or isinstance(data, np.uint8):
        return format(data, "08b")
    else:
        raise TypeError("Type not supported.")
        
# Creates an image with the RGB values
def createImage(data, resolution):
    image = Image.new("RGB", resolution)
    image.putdata(data)
    return image


def remove_n_least_significant_bits(value, n):
    value = value >> n
    return value << n


def get_n_least_significant_bits(value, n):
    value = value << MAX_BIT_VALUE - n
    value = value % MAX_COLOR_VALUE
    return value >> MAX_BIT_VALUE - n


def get_n_most_significant_bits(value, n):
    return value >> MAX_BIT_VALUE - n


def shift_n_bits_to_8(value, n):
    return value << MAX_BIT_VALUE - n


# Function to encode the image that we want to hide, into the cover image
def encodeTextFile(file_to_hide, image_to_hide_in, n_bits):
    width = image_to_hide_in.size[0]
    height = image_to_hide_in.size[1]
    hide_in_image = image_to_hide_in.load()


    # List to store the ascii values from the content of the file
    file_bytes = []

    #Uncomment to encode txt file in img
    with open(file_to_hide, "rb") as out:
        while True:
            c = out.read(1)
            #print(c)
            if not c:
                print("End of file")
                break
            file_bytes.append(ord(c))

    #Uncomment to encode docx file in img
    """fullText = []
    doc = docx.Document(file_to_hide)
    for i in doc.paragraphs:
        #print(i.text)
        fullText.append(i.text)
    docText = '\n'.join(fullText)
    for i in docText:
        #print(ord(i))
        file_bytes.append(ord(i))"""


    # Check size of cover img
    img_size = (width * height) * 3
    print("The img size is ", img_size)

    # Keep adding character till it can be reshaped into a list with 3 elements (In the form of RGB (1,2,3))
    while (len(file_bytes) % 3 != 0):
        file_bytes.append(ord("A"))
        # print(file_size % 3)

    file_bytes = np.array(file_bytes)
    file_bytes = np.reshape(file_bytes, (-1, 3))
    print(len(file_bytes))
    """for i in file_bytes:
        print(i)"""

    file_size = os.path.getsize(file_to_hide)  # file_size is the same as len(file_bytes)
    print("The payload file size is ", file_size)

    # Check if cover file size is smaller than payload img size
    if (img_size < file_size):
        print("Cover object size is smaller than payload size")
        quit()

    # this will store the values of each individual pixel as a matrix.
    data = []
    counter = 0  # Loop thru the txt file
    # looping the hide_image object.
    print("HEIGHT ", height)
    print("WIDTH ", width)
    for y in range(height):
        
        for x in range(width):
            
            # print(counter)

            # most sig bits
            # print(hide_image[x,y]) #Uncomment this to see the pixel values in r,g,b form. (Shows every pixel's RGB value)
            try:
                # the value of n can be 1 or 2 and you won't see much difference in the encoded image.
                # gets n most significant bits of r,g,b values of image to hide.
                # print("HIDE IMAGE ", hide_image[x,y])
                if (counter < len(file_bytes)):
                    r_hide, g_hide, b_hide = file_bytes[counter]
                    #print(r_hide," ", g_hide," ", b_hide)

                    r_hide = get_n_most_significant_bits(r_hide, n_bits)
                    g_hide = get_n_most_significant_bits(g_hide, n_bits)
                    b_hide = get_n_most_significant_bits(b_hide, n_bits)
                    

                # remove least n significant bits of txt file's content to hide in so we can store
                # the n most significant bits in that place.

                r_hide_in, g_hide_in, b_hide_in = hide_in_image[x,y]

                r_hide_in = remove_n_least_significant_bits(r_hide_in, n_bits)
                g_hide_in = remove_n_least_significant_bits(g_hide_in, n_bits)
                b_hide_in = remove_n_least_significant_bits(b_hide_in, n_bits)
                

                # Adding each MSB R,G,B from payload to LSB R,G,B cover image
                if (counter >= len(file_bytes)):
                    data.append((0 + r_hide_in,
                                0 + g_hide_in,
                                0 + b_hide_in))
                else:
                    data.append((r_hide + r_hide_in,
                                g_hide + g_hide_in,
                                b_hide + b_hide_in))

                counter += 1
                

            # incase of exception it will show the reason.
            except Exception as e:
                print(e)

    #return an Image object from the above data.
    return createImage(data, image_to_hide_in.size)

#Decode image and n_bits as parameters.
def decodeImage(image_to_decode, n_bits):
    width = image_to_decode.size[0]
    height = image_to_decode.size[1]
    encoded_image = image_to_decode.load()

    #matrix that will store the extracted pixel values from the encoded Image.
    data = []

    counter = 0
    
    #looping through every pixel in the encoded Image.
    for y in range(height):
        for x in range(width):

            #gets rgb values of encoded image.
            r_encoded, g_encoded, b_encoded = encoded_image[x,y]

            #get n least significant bits for each r,g,b value of the encoded image
            r_encoded = get_n_least_significant_bits(r_encoded, n_bits)
            g_encoded = get_n_least_significant_bits(g_encoded, n_bits)
            b_encoded = get_n_least_significant_bits(b_encoded, n_bits)

            #shifts the n bits to right so that they occupy a total of 8 bit spaces.
            #If 10 are the bits then shifting them would look like 10000000
            #this would ofcourse be converted to an int as per python's bit operations.
            
            r_encoded = shift_n_bits_to_8(r_encoded, n_bits)
            g_encoded = shift_n_bits_to_8(g_encoded, n_bits)
            b_encoded = shift_n_bits_to_8(b_encoded, n_bits)

            data.append((r_encoded, g_encoded, b_encoded))
            counter += 1
            """if (counter == 12666):
                break"""
            
    data = np.array(data)
    data = data.flatten()
    print(data)

    #Uncomment to retrieve txt file from image
    with open("decoded_file.txt", "w", encoding="utf-8") as f:
        for i in data:
            f.write(chr(i))

    #Uncomment to retrieve docx file from image
    """textList = []
    decoded_doc = docx.Document()
    #with open("decoded_docx.docx", "w", encoding="utf-8") as f:
    for i in data:
        #print(chr(i))
        x = chr(i)
        #print(x)
        cleaned_string = ''.join(c for c in x if valid_xml_char_ordinal(c))
        #print(cleaned_string)
        textList.append(cleaned_string)
    
    fullText = ""
    for i in textList:
        fullText += i
    #print(fullText)
    decoded_doc.add_paragraph(fullText)
    decoded_doc.save("decoded_docx.docx")"""

    
    
    return None
    return createImage(data, image_to_decode.size)

# Decode image and n_bits as parameters.
def decodeFile(file_to_decode, n_bits, height, width):
    
    # List to store the ascii values from the content of the file
    asciiArray = []
    with open(file_to_decode, "r", encoding="utf-8") as f:
        for i in f.read():
            asciiArray.append(ord(i))
    # Keep adding character till it can be reshaped into a list with 3 elements (In the form of RGB (1,2,3))
    while (len(asciiArray) % 3 != 0):
        asciiArray.append(ord("A"))
        # print(file_size % 3)
    npAsciiArray = np.array(asciiArray)
    npAsciiArray = np.reshape(npAsciiArray, (-1, 3))
    # print("ASCIITEST ", npAsciiArray)

    # matrix that will store the extracted pixel values from the encoded txt file.
    data = []
    counter = 0
    # looping through every pixel in the encoded Image.
    for y in range(height):
        for x in range(width):
            # gets rgb values of encoded txt file.
            r_encoded, g_encoded, b_encoded = npAsciiArray[counter]

            # get n least significant bits for each r,g,b value of the encoded image
            r_encoded = get_n_least_significant_bits(r_encoded, n_bits)
            g_encoded = get_n_least_significant_bits(g_encoded, n_bits)
            b_encoded = get_n_least_significant_bits(b_encoded, n_bits)

            # shifts the n bits to right so that they occupy a total of 8 bit spaces.
            # If 10 are the bits then shifting them would look like 10000000
            # this would ofcourse be converted to an int as per python's bit operations.

            r_encoded = shift_n_bits_to_8(r_encoded, n_bits)
            g_encoded = shift_n_bits_to_8(g_encoded, n_bits)
            b_encoded = shift_n_bits_to_8(b_encoded, n_bits)

            data.append((r_encoded, g_encoded, b_encoded))
            counter += 1
    # print("DATA ", data)

    return createImage(data, (width, height))


# # Running encode function
# # no. of bits
n_bits = 8
# # path of files replace them as per your need.
#image_to_hide_path = "tree2.jpg"
#image_to_hide_in_path = "test.txt"
#image_to_hide_in_path = "wordTest.docx"
#image_to_hide_in_path = "excel.xlsx"
file_to_hide_path = "test.txt"
cover_img_path = "pics/tree.jpg"
image_to_hide = Image.open(cover_img_path)
#height, width = encodeImage(image_to_hide, image_to_hide_in_path, n_bits)
encoded_image_path = "pics/encoded_img_from_file.png"
encodeTextFile(file_to_hide_path, image_to_hide, n_bits).save(encoded_image_path)
print("File encoded Successfully!")
#
# # running the decode function
n_bits = 8
#
# # path where you would want to save decoded Image.
image_to_decode_path = "pics/encoded_img_from_docx.png"
image_to_decode = Image.open(image_to_decode_path)
decodeImage(image_to_decode, n_bits)
print("Image decoded Successfully!")
